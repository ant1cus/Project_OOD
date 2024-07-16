# import datetime
import os
import pathlib
import re
import threading

import shutil
import time
import traceback
from docx.shared import Pt

import docx
from PyQt5.QtCore import QThread, pyqtSignal
from DoingWindow import CheckWindow


class CancelException(Exception):
    pass


class ChangeNumberInstance(QThread):
    status_finish = pyqtSignal(str, str)
    progress_value = pyqtSignal(int)
    info_value = pyqtSignal(str, str)
    status = pyqtSignal(str)
    line_progress = pyqtSignal(str)
    line_doing = pyqtSignal(str)

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path_old = incoming_data['path_old']
        self.path_new = incoming_data['path_new']
        self.set_number = incoming_data['set_number']
        self.logging = incoming_data['logging']
        self.queue = incoming_data['queue']
        self.default_path = incoming_data['default_path']
        self.event = threading.Event()
        self.event.set()
        self.all_doc = 0
        self.now_doc = 0
        self.percent_progress = 0
        self.move = incoming_data['move']
        self.name_dir = pathlib.Path(self.path_old).name
        title = f'Изменение номера экземпляра в папке «{self.name_dir}»'
        self.window_check = CheckWindow(self.default_path, self.event, self.move, title)
        self.progress_value.connect(self.window_check.progressBar.setValue)
        self.line_progress.connect(self.window_check.lineEdit_progress.setText)
        self.line_doing.connect(self.window_check.lineEdit_doing.setText)
        self.info_value.connect(self.window_check.info_message)
        self.window_check.show()

    def run(self):
        try:
            current_progress = 0
            self.logging.info('Начинаем менять номера комплектов')
            self.line_progress.emit(f'Выполнено {int(current_progress)} %')
            self.progress_value.emit(0)
            self.all_doc = len(os.listdir(self.path_old))*len(self.set_number)
            self.percent_progress = 100/(len(os.listdir(self.path_old))*len(self.set_number))
            for number_folder in self.set_number:
                os.makedirs(pathlib.Path(self.path_new, str(number_folder) + ' экземпляр'), exist_ok=True)
                for doc in [file for file in os.listdir(self.path_old) if file.endswith('.docx')]:
                    self.now_doc += 1
                    self.line_doing.emit(f'Создаем {str(number_folder)} экземпляр для {doc}'
                                         f' ({self.now_doc} из {self.all_doc})')
                    self.logging.info(f'Создаем {str(number_folder)} экземпляр для {doc}'
                                      f' ({self.now_doc} из {self.all_doc})')
                    self.event.wait()
                    if self.window_check.stop_threading:
                        raise CancelException()
                    shutil.copy2(pathlib.Path(self.path_old, doc),
                                 pathlib.Path(self.path_new, str(number_folder) + ' экземпляр'))
                    doc_2 = docx.Document(pathlib.Path(self.path_new, str(number_folder) + ' экземпляр', doc))
                    for p_2 in doc_2.sections[0].first_page_header.paragraphs:
                        if re.findall(r'№1', p_2.text):
                            text = re.sub(r'№1', '№' + str(number_folder), p_2.text)
                            p_2.text = text
                            for run in p_2.runs:
                                run.font.size = Pt(11)
                                run.font.name = 'Times New Roman'
                            break
                    for p_2 in doc_2.sections[len(doc_2.sections) - 1].first_page_footer.paragraphs:
                        if re.findall(r'Отп. 1 экз. в адрес', p_2.text):
                            text = re.sub(r'Отп. 1 экз. в адрес', 'Отп. ' + str(number_folder) + ' экз. в адрес',
                                          p_2.text)
                            p_2.text = text
                            for run in p_2.runs:
                                run.font.size = Pt(11)
                                run.font.name = 'Times New Roman'
                            break
                    doc_2.save(pathlib.Path(self.path_new, str(number_folder) + ' экземпляр', doc))  # Сохраняем
                    current_progress += self.percent_progress
                    self.line_progress.emit(f'Выполнено {int(current_progress)} %')
                    self.progress_value.emit(int(current_progress))
            self.line_progress.emit(f'Выполнено 100 %')
            self.logging.info(f"Создание экземпляров документов в папке «{self.name_dir}» успешно завершено")
            self.progress_value.emit(int(100))
            os.chdir(self.default_path)
            self.status.emit(f"Создание экземпляров документов в папке «{self.name_dir}» успешно завершено")
            self.status_finish.emit('change_number_instance', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            # print(datetime.datetime.now() - start_time)
            return
        except CancelException:
            self.logging.warning(f"Создание экземпляров документов в папке «{self.name_dir}» отменено пользователем")
            self.status.emit(f"Создание экземпляров документов в папке «{self.name_dir}» отменено пользователем")
            os.chdir(self.default_path)
            self.status_finish.emit('change_number_instance', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
        except BaseException as es:
            self.logging.error(es)
            self.logging.error(traceback.format_exc())
            self.logging.warning(f"Создание экземпляров документов в папке «{self.name_dir}» не заврешено из-за ошибки")
            self.info_value.emit('УПС!', 'Работа программы завершена из-за непредвиденной ошибки')
            self.event.clear()
            self.event.wait()
            self.status.emit(f"Ошибка при создании экземпляров документов в папке «{self.name_dir}»")
            os.chdir(self.default_path)
            self.status_finish.emit('change_number_instance', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return

