import os
import docx
import openpyxl
import re
import traceback
from natsort import natsorted
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PyQt5.QtCore import QThread, pyqtSignal
import win32com
from win32com.client import Dispatch


class ZoneChecked(QThread):
    progress = pyqtSignal(int)  # Сигнал для прогресс бара
    status = pyqtSignal(str)  # Сигнал для статус бара

    def __init__(self, output):  # Список переданных элементов.
        QThread.__init__(self)
        self.path = output[0]
        self.table = output[1]
        self.department = output[2]
        self.win_lin = output[3]
        self.zone = output[4]
        self.one_table = output[5]
        self.logging = output[6]

    def run(self):
        progress = 0
        self.logging.info("Начинаем")
        self.status.emit('Старт')
        self.progress.emit(progress)

        zone_name = ('Стац.', 'Воз.', 'Нос.', 'r1', 'r1`')
        os.chdir(self.path)
        self.logging.info("Сортировка")
        docs = [i for i in os.listdir('.') if i[-4:] == 'docx']
        docs = natsorted(docs)
        percent = 100 / len(docs)
        void = 0
        e = []
        self.logging.info("Проходимся по списку")
        try:
            for i in docs:
                self.logging.info("Документ " + str(i) + " в работе")
                self.status.emit('Проверяем документ ' + i)
                if '~' not in i:
                    # Для того, что бы не съезжала заливка ее нужно добавлять каждый раз для каждой ячейки
                    shading_elm = []
                    string = {}  # Для записи непроходящих частот
                    shading_index = 0  # Счетчик для заливки
                    doc = docx.Document(i)
                    table = doc.tables[int(self.table) - 1]  # Таблица для проверки (общая)
                    if void == 1:
                        e.append('\n')
                        void = 2
                    errors = [i.partition(' ')[2][:-5]]
                    win_lin = 10 if self.win_lin else 5  # Если 2 системы
                    zone = 0
                    self.logging.info("Считываем зоны")
                    if self.department:  # Если ФСБ
                        for j in range(0, win_lin, 1):
                            if j == 0 and self.win_lin:
                                errors.append('\n')
                                errors.append('Windows')
                                n_s = table.cell(2, 1).text  # Чтобы проверять и не считывать другую систему
                                name_system = re.findall(r'^\w+\b\s(\b\w+\b)', n_s)[0]
                            if j == 5 and self.win_lin:
                                if len(string) != 0:  # Добавляем частоты если они есть.
                                    errors.append(string)
                                    string = {}
                                errors.append('\n')
                                errors.append('Linux')
                                n_s = table.cell(6, 1).text  # Чтобы проверять имя системы
                                name_system = re.findall(r'^\w+\b\s(\b\w+\b\s\b\w+\b)', n_s)[0]
                            if j <= 2:
                                if self.win_lin:
                                    zone = table.cell(3, j + 2).text.replace(',', '.')
                                else:
                                    zone = table.cell(2, j + 2).text.replace(',', '.')
                            elif j <= 4:
                                if self.win_lin:
                                    zone = table.cell(j + 1, 2).text.replace(',', '.')
                                else:
                                    zone = table.cell(j, 2).text.replace(',', '.')
                            elif j <= 7 and self.win_lin:
                                zone = table.cell(7, j - 3).text.replace(',', '.')
                            elif j <= 9 and self.win_lin:
                                zone = table.cell(j, 2).text.replace(',', '.')
                            try:
                                if void == 0:  # Добавляем имена диапазонов
                                    e.append('\t')
                                    for k in zone_name:
                                        e.append(k)
                                    void = 2
                                    e.append('\n')
                                void = 1
                                errors.append(round(float(zone), 1))
                                if self.one_table is False:  # Если нужно проверять таблицу
                                    self.logging.info("Проверяем и красим таблицу")
                                    self.status.emit('Проверяем и закрашиваем таблицу в документе ' + str(i))
                                    percent_ = percent/5
                                    if j in self.zone:
                                        # Условия проверки
                                        if (float(self.zone[j]) < float(zone)) and (float(self.zone[j]) != 0):
                                            flag_for_exit = 1  # Для прерывания цикла
                                            # Поисковый диапазон
                                            range_search = [9, 11, 13, 5, 7] * 2 if self.win_lin else [9, 11, 13, 5, 7]
                                            x = 0
                                            name = ''
                                            table_3 = doc.tables[1] if self.win_lin else doc.tables[2]
                                            for row in table_3.rows:
                                                for cell in row.cells:
                                                    try:
                                                        # Для определения позиции ячейки (tc.top, tc.bottom etc)
                                                        tc = cell._tc
                                                        # Если полностью объединена
                                                        if tc.right - tc.left == len(table_3.columns):
                                                            if cell.text != 'Опасные сигналы не обнаружены':
                                                                if self.win_lin:
                                                                    # Если находим имя - не прерываем цикл
                                                                    if len(re.findall(name_system, cell.text)):
                                                                        flag_for_exit = 1
                                                                    else:
                                                                        flag_for_exit = 0
                                                                # Имя системы
                                                                name = re.findall(r"\(([^)]*)\)", cell.text)[0]
                                                        if flag_for_exit:  # Если нужно в цикл
                                                            if tc.right == 1 and tc.left == 0:
                                                                frq = float(cell.text.replace(',', '.'))  # Частота
                                                                try:
                                                                    x = float(table_3.cell(tc.top, range_search[j])
                                                                              .text.replace(',', '.'))
                                                                except BaseException:
                                                                    if '<' in table_3.cell(tc.top, range_search[j]).text:
                                                                        x = -1
                                                                # Если больше, то красим через xml
                                                                if x > float(self.zone[j]):
                                                                    string.setdefault(name, [])
                                                                    shading_elm.append(parse_xml(
                                                                        r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                                                                    )
                                                                    table_3.rows[tc.top].cells[0]._tc.get_or_add_tcPr().append(
                                                                        shading_elm[shading_index])
                                                                    shading_index += 1
                                                                    shading_elm.append(parse_xml(
                                                                        r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                                                                    )
                                                                    table_3.rows[tc.top].cells[
                                                                        range_search[j]]._tc.get_or_add_tcPr().append(
                                                                        shading_elm[shading_index])
                                                                    shading_index += 1
                                                                    if frq not in string[name]:
                                                                        string[name].append(frq)
                                                        break
                                                    except BaseException:
                                                        break
                                            self.logging.info("Сохраняем документ")
                                            doc.save(os.path.abspath(self.path) + '\\' + i)
                                    progress = progress + percent_
                                    self.progress.emit(progress)
                            except ValueError:
                                errors.append(zone)
                    else:
                        self.logging.info("Считываем зоны")
                        for j in range(0, 4, 1):
                            zone = table.cell(j + 1, 1).text.replace(',', '.')
                            try:
                                if void == 0:
                                    e.append('\t')
                                    for k in zone_name[:-1]:
                                        e.append(k)
                                    void = 2
                                    e.append('\n')
                                void = 1
                                try:
                                    errors.append(int(zone))
                                except ValueError:
                                    errors.append(round(float(zone), 1))
                                if self.one_table is False:
                                    self.logging.info("Проверяем и красим таблицу")
                                    self.status.emit('Проверяем и закрашиваем таблицу в документе ' + str(i))
                                    percent_ = percent / 4
                                    if j in self.zone:
                                        # Условия проверки
                                        if (float(self.zone[j]) < float(zone)) and (float(self.zone[j]) != 0):
                                            flag_for_exit = 0
                                            range_search = [7, 8, 9, 10]
                                            x = 0
                                            table_3 = doc.tables[2]
                                            for row in table_3.rows:
                                                for cell in row.cells:
                                                    try:
                                                        tc = cell._tc
                                                        if tc.right - tc.left == len(table_3.columns):
                                                            if cell.text == '3 категория':
                                                                flag_for_exit = 1
                                                            try:
                                                                try:
                                                                    x = int(table_3.cell(tc.top, range_search[j]).text)
                                                                except ValueError:
                                                                    x = float(table_3.cell(tc.top, range_search[j])
                                                                              .text.replace(',', '.'))
                                                            except BaseException:
                                                                if '<' in table_3.cell(tc.top, range_search[j]).text:
                                                                    x = -1
                                                            if x > float(self.zone[j]):
                                                                # string.setdefault(name, [])
                                                                shading_elm.append(parse_xml(
                                                                    r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                                                                table_3.rows[tc.top].cells[1]._tc.get_or_add_tcPr().append(
                                                                    shading_elm[shading_index])
                                                                shading_index += 1
                                                                shading_elm.append(parse_xml(
                                                                    r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                                                                table_3.rows[tc.top].cells[2]._tc.get_or_add_tcPr().append(
                                                                    shading_elm[shading_index])
                                                                shading_index += 1
                                                                shading_elm.append(parse_xml(
                                                                    r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                                                                table_3.rows[tc.top].cells[
                                                                    range_search[j]]._tc.get_or_add_tcPr().append(
                                                                    shading_elm[shading_index])
                                                                shading_index += 1
                                                        break
                                                    except BaseException:
                                                        break
                                                if flag_for_exit:
                                                    break
                                            self.logging.info("Сохраняем документ")
                                            doc.save(os.path.abspath(self.path) + '\\' + i)
                                    progress = progress + percent_
                                    self.progress.emit(progress)
                            except ValueError:
                                pass
                    self.logging.info("Добавляем результаты")
                    self.status.emit('Добавляем результаты документа ' + str(i))
                    if void == 1:
                        for el in errors:
                            if type(el) != dict:
                                e.append(str(el))
                            else:
                                e.append(el)
                        if len(string) > 0:
                            e.append(string)
        except BaseException as es:
            with open('log.txt', mode='w') as f:
                print('--------------------------------------', file=f)
                print(es, file=f)
                print('--------------------------------------', file=f)
                print(traceback.format_exc(), file=f)
            progress = 0
            self.progress.emit(progress)
            os.chdir('C://')
            if es.args[2][2] == 'Запрашиваемый номер семейства не существует.':
                self.status.emit('Не верно указан номер таблицы')
            return
        e.append('\n')
        self.logging.info("Формируем excel")
        self.status.emit('Формируем отчёт')
        if self.department:
            zone = [self.zone.get(i) for i in range(0, 5)]
            len_fill = len(zone_name)
        else:
            zone = [self.zone.get(i) for i in range(0, 4)]
            len_fill = len(zone_name[:-1])
        thin = openpyxl.styles.Side(border_style="thin", color="000000")
        wb = openpyxl.Workbook()
        ws = wb.active
        i, j = 1, 1
        for el in e:
            if i != 1 and el != 'Windows' and el != 'Linux':  # Чтобы отдельно заполнять первый столбец
                if el == '\n':  # Если новая строка (используется как разделитель)
                    pass
                else:
                    if j != 1:  # Если столбец не первый
                        if '<' in el:  # Если значения '<0.1'
                            ws.cell(i, j).value = el  # Просто пишем
                        else:
                            if type(el) == dict:  # Если тип словарь (если не прошло и есть значения для записи)
                                j_ = 0
                                for key in el:  # Имя и значения. Форматирование.
                                    ws.cell(i, j + j_).value = key
                                    ws.cell(i, j + j_).alignment = openpyxl.styles.Alignment(horizontal="left",
                                                                                             vertical="center")
                                    ws.cell(i, j + j_).font = openpyxl.styles.Font(bold=True, name="Times New Roman",
                                                                                   size="11")
                                    j_ += 1
                                    for element in sorted(el[key]):
                                        ws.cell(i, j + j_).value = element
                                        ws.cell(i, j + j_).alignment = openpyxl.styles.Alignment(horizontal="left",
                                                                                                 vertical="center")
                                        ws.cell(i, j + j_).font = openpyxl.styles.Font(name="Times New Roman",
                                                                                       size="11")
                                        j_ += 1
                            elif '.' in el:  # Если есть точка, то форматируем
                                ws.cell(i, j).number_format = '0.0'
                                ws.cell(i, j).value = float(el)
                            else:  # Если целочисленные (для ФСТЭК в основном)
                                ws.cell(i, j).value = int(el)
                    else:  # Если столбец первый
                        len_ = len(el.partition('.')[2])  # Определяем какая длина после запятой (номер комплекта)
                        ws.cell(i, j).number_format = '0.' + '0' * len_  # Для форматирования строки под требования
                        f = '.' + str(len_) + 'f'
                        ws.cell(i, j).value = float(format(float(el), f))  # Забиваем значение
                    if j < len_fill + 2:  # Форматируем ячейки
                        if j == 1:
                            ws.cell(i, j).alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")
                        else:
                            ws.cell(i, j).border = openpyxl.styles.Border(top=thin, left=thin, right=thin, bottom=thin)
                            ws.cell(i, j).alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
                        ws.cell(i, j).font = openpyxl.styles.Font(name="Times New Roman", size="11")
                        ws.cell(i, j).fill = openpyxl.styles.PatternFill(start_color='92D050',
                                                                         end_color='92D050',
                                                                         fill_type="solid")
            else:  # Форматируем и вставляем заголовки.
                ws.cell(i, j).value = el
                if i != 1:
                    ws.cell(i, j).alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")
                    ws.cell(i, j).fill = openpyxl.styles.PatternFill(start_color='92D050',
                                                                     end_color='92D050',
                                                                     fill_type="solid")
                else:
                    ws.cell(i, j).alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
                ws.cell(i, j).font = openpyxl.styles.Font(name="Times New Roman", size="11")
            j += 1  # Увеличиваем столбец
            if el == '\n':  # Если переход к новой строке.
                flag = 0
                for element in range(2, len_fill + 2):  # Проверяем на проходимость
                    if i != 1 and type(ws.cell(i, element).value) != str and ws.cell(i, element).value:
                        try:  # Если не проходит, то форматируем
                            if (ws.cell(i, element).value > float(zone[element - 2])) and (float(zone[element - 2]) != 0):
                                ws.cell(i, element).font = openpyxl.styles.Font(bold=True, name="Times New Roman")
                                flag = 1
                        except TypeError:
                            pass
                if flag:  # Если что-то форматируется, то закрашиваем
                    if self.win_lin:
                        for min_ in range(1, 3):  # Закрашиваем номер комплекта
                            print(ws.cell(i - min_, 1).value)
                            if type(ws.cell(i - min_, 1).value) == float:
                                ws.cell(i - min_, 1).alignment = openpyxl.styles.Alignment(
                                    horizontal="left", vertical="center")
                                ws.cell(i - min_, 1).fill = openpyxl.styles.PatternFill(
                                    start_color='FFC000', end_color='FFC000', fill_type="solid")
                                break
                    for element in range(1, len_fill + 2):  # Закрашиваем строки
                        ws.cell(i, element).fill = openpyxl.styles.PatternFill(
                            start_color='FFC000', end_color='FFC000', fill_type="solid")
                i += 1
                j = 1
        wb.save(filename='Зоны.xlsx')  # Сохраняем книгу.
        wb.close()
        os.startfile('Зоны.xlsx')
        progress = 100
        self.progress.emit(progress)
        self.status.emit('Готово!')
        os.chdir('C://')
