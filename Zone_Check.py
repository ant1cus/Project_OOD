# import datetime
import os
import pathlib
import threading
import time
import docx
import re
import traceback
import pandas as pd
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from itertools import groupby
from openpyxl import load_workbook
from openpyxl.styles import Alignment, PatternFill, Font, Side, Border
from natsort import natsorted
from PyQt5.QtCore import QThread, pyqtSignal
from DoingWindow import CheckWindow


def set_interval(interval):
    def decorator(function):
        def wrapper(*args, **kwargs):
            stopped = threading.Event()

            def loop():  # executed in another thread
                while not stopped.wait(interval):  # until stopped
                    function(*args, **kwargs)

            t = threading.Thread(target=loop)
            t.daemon = True  # stop if the program exits
            t.start()
            return stopped
        return wrapper
    return decorator


class CancelException(Exception):
    pass


class ZoneChecked(QThread):
    status_finish = pyqtSignal(str, str)
    progress_value = pyqtSignal(int)
    info_value = pyqtSignal(str, str)
    status = pyqtSignal(str)
    line_progress = pyqtSignal(str)
    line_doing = pyqtSignal(str)

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path = incoming_data['path_check']
        self.table = incoming_data['table_number']
        self.department = incoming_data['department']
        self.win_lin = incoming_data['win_lin']
        self.extend_report = incoming_data['extend_report']
        self.zone = incoming_data['zone_all']
        self.one_table = incoming_data['one_table']
        self.logging = incoming_data['logging']
        self.queue = incoming_data['queue']
        self.default_path = incoming_data['default_path']
        self.event = threading.Event()
        self.event.set()
        self.move = incoming_data['move']
        self.name_dir = pathlib.Path(self.path).name
        title = f'Проверка зон в папке «{self.name_dir}»'
        self.window_check = CheckWindow(self.default_path, self.event, self.move, title)
        self.progress_value.connect(self.window_check.progressBar.setValue)
        self.line_progress.connect(self.window_check.lineEdit_progress.setText)
        self.line_doing.connect(self.window_check.lineEdit_doing.setText)
        self.info_value.connect(self.window_check.info_message)
        self.app_text = ''
        self.previous_text = ''
        self.window_check.show()
        self.timer_line_progress()

    @set_interval(3)
    def timer_line_progress(self):
        text = self.window_check.lineEdit_progress.text()
        if '•' in text:
            text = re.sub('•', '', text)
        if text == self.previous_text:
            self.app_text = self.app_text + '•' if len(self.app_text) <= 5 else ''
            self.line_progress.emit(text + self.app_text)
        # time.sleep(1)

    def set_line_progress(self, text):
        self.previous_text = text
        self.line_progress.emit(text)
        self.app_text = ''

    def run(self):

        def round_set(ws_cell) -> None:
            ws_cell.number_format = '0.' + '0' * len(str(ws_cell.value).partition('.')[2])
            ws_cell.value = float(format(float(ws_cell.value), '.' + str(len(str(ws_cell.value).partition('.')[2]))
                                         + 'f'))
            return

        def round_number(ws_cell, val='', bold=False, align=False, string=False) -> None:
            if not val:
                val = ws_cell.value
            align_val = "center" if align else "left"
            if string:
                ws_cell.value = val
            else:
                len_val = len(str(val).partition('.')[2])
                if self.department or len_val > 0:
                    if self.department and len_val == 0:
                        len_val = 1
                    ws_cell.number_format = '0.' + '0' * len_val
                    ws_cell.value = float(format(float(val), f'.{len_val}f'))
                else:
                    ws_cell.number_format = '0'
                    ws_cell.value = int(val)
            ws_cell.alignment = Alignment(horizontal=align_val, vertical="center")
            ws_cell.font = Font(bold=bold, name="Times New Roman", size="11")
            return

        self.progress_value.emit(0)
        progress = 0
        report = {}
        self.logging.info('Начинаем проверять зоны')
        self.logging.info("Сортировка")
        docs = [i for i in os.listdir(self.path) if i[-4:] == 'docx' and '~' not in i]
        docs = natsorted(docs)
        try:
            percent = 100 / (len(docs) * 2)
            all_doc = len(docs)
            now_doc = 0
            self.logging.info("Проходимся по списку")
            zone_label = ['Стац.', 'Воз.', 'Нос.', 'r1', 'r1`'] if self.department else ['Стац.', 'Воз.', 'Нос.', 'r1']
            df_excel = pd.DataFrame(data=None, columns=['№'] + zone_label)
            df_extended_report = pd.DataFrame(data=None, columns=['№'] + zone_label)
            self.set_line_progress(f'Выполнено {int(progress)} %')
            # self.line_progress.emit(f'Выполнено {int(progress)} %')
            # start_time = datetime.datetime.now()
            for name_doc in docs:
                now_doc += 1
                self.logging.info("Документ " + str(name_doc) + " в работе")
                self.line_doing.emit(f'Проверяем {str(name_doc)} ({now_doc} из {all_doc})')
                self.event.wait()
                if self.window_check.stop_threading:
                    raise CancelException()
                set_number = name_doc.rpartition(' ')[2][:-5]
                report[set_number] = {'error_win': False, 'error_lin': False, 'win': {}, 'lin': {}}
                # Для того, что бы не съезжала заливка ее нужно добавлять каждый раз для каждой ячейки
                doc = docx.Document(pathlib.Path(self.path, name_doc))
                table_zone = doc.tables[int(self.table) - 1]  # Таблица для проверки (общая)
                self.logging.info("Считываем зоны")
                # Попытка по новому прочитать и ускорить процесс
                user_zone = [self.zone[3], self.zone[4], self.zone[0], self.zone[1], self.zone[2]] if self.department\
                    else [self.zone[i] for i in range(4)]
                user_zone = [round(float(i), 1) for i in user_zone]
                line_for_append = [set_number] + ['']*len(zone_label)
                df_excel.loc[len(df_excel)] = line_for_append
                df_win_index = len(df_excel) - 1
                df_lin_index = 0
                if self.win_lin:
                    line_for_append[0] = 'Windows'
                    df_excel.loc[len(df_excel)] = line_for_append
                    df_win_index = len(df_excel) - 1
                    line_for_append[0] = 'Linux'
                    df_excel.loc[len(df_excel)] = line_for_append
                    df_lin_index = len(df_excel) - 1
                df_list = []
                # Собираем все таблицы и делаем из них DataFrame
                table_value = doc.tables[1] if self.win_lin and self.department else doc.tables[2]
                table_pwr = doc.tables[3] if self.win_lin and self.department else doc.tables[4]
                self.line_doing.emit(f'Считываем зоны в {str(name_doc)} ({now_doc} из {all_doc})')
                table_for_check = [table_zone] if self.one_table else [table_zone, table_value, table_pwr]
                for table_val in table_for_check:
                    df = pd.DataFrame()
                    for index in range(len(table_val.columns)):
                        df[index] = list(map(lambda val: '0' if '<' in val.text else val.text.replace(',', '.'),
                                             table_val.column_cells(index)))
                    df_list.append(df)
                if self.one_table is False:
                    for i in [1, 2]:
                        cat_3 = df_list[i].loc[df_list[i][0] == '3 категория'].index.to_list()
                        if cat_3 and cat_3[0]:
                            df_list[i] = df_list[i].drop(index=[i for i in range(cat_3[0], df_list[i].shape[0])])
                # Во всех таблицах удаляем первые 2 строки, а в первой дополнительно удаляем первые 2 колонки
                drop_line = [0, 1] if self.department else [0]
                df_list[0] = df_list[0].drop(columns=drop_line, axis=1)
                for i in range(3 if self.one_table is False else 1):
                    df_list[i] = df_list[i].drop(index=drop_line)
                    df_list[i].reset_index(drop=True, inplace=True)
                # Если в 1-ой таблице есть объединенная строка (имя системы) - значит будут 2 системы
                # Сразу добавляем строки в таблицу для вывода в excel
                # Вносим в список зоны для проверки
                zone_check = {}
                name_win = 'default'
                name_lin = 'Linux'
                if self.department:
                    if all(x == df_list[0].iloc[0, 0] for x in df_list[0].iloc[0]):
                        name_win = 'Windows'
                else:
                    if df_list[0].shape[0] > 5:
                        name_win = 'Windows'
                if name_win == 'default':
                    if self.department:
                        zone_check[name_win] = [df_list[0].iloc[1, 0], df_list[0].iloc[2, 0]]
                        for enum, x in enumerate(['Стац.', 'Воз.', 'Нос.']):
                            zone_check[name_win].append(df_list[0].iloc[0, enum])
                            df_excel.loc[df_win_index, x] = df_list[0].iloc[0, enum]
                        df_excel.loc[df_win_index, 'r1'] = df_list[0].iloc[1, 0]
                        df_excel.loc[df_win_index, 'r1`'] = df_list[0].iloc[2, 0]
                        zone_check[name_lin] = []
                    else:
                        zone_check[name_win] = [df_list[0].iloc[i, 0] for i in range(4)]
                        for enum, x in enumerate(zone_label):
                            df_excel.loc[df_win_index, x] = df_list[0].iloc[enum, 0]
                        zone_check[name_lin] = []
                else:
                    if self.department:
                        zone_check[name_win] = [df_list[0].iloc[2, 0], df_list[0].iloc[3, 0]]
                        for enum, x in enumerate(['Стац.', 'Воз.', 'Нос.']):
                            zone_check[name_win].append(df_list[0].iloc[1, enum])
                            df_excel.loc[df_win_index, x] = df_list[0].iloc[1, enum]
                        df_excel.loc[df_win_index, 'r1'] = df_list[0].iloc[2, 0]
                        df_excel.loc[df_win_index, 'r1`'] = df_list[0].iloc[3, 0]
                        zone_check[name_lin] = [df_list[0].iloc[6, 0], df_list[0].iloc[7, 0]]
                        for enum, x in enumerate(['Стац.', 'Воз.', 'Нос.']):
                            zone_check[name_lin].append(df_list[0].iloc[5, enum])
                            df_excel.loc[df_lin_index, x] = df_list[0].iloc[5, enum]
                        df_excel.loc[df_lin_index, 'r1'] = df_list[0].iloc[6, 0]
                        df_excel.loc[df_lin_index, 'r1`'] = df_list[0].iloc[7, 0]
                    else:
                        zone_check[name_win] = [df_list[0].iloc[i, 0] for i in range(1, 5)]
                        for enum, x in enumerate(zone_label):
                            df_excel.loc[df_win_index, x] = df_list[0].iloc[enum + 1, 0]
                        zone_check[name_lin] = [df_list[0].iloc[i, 0] for i in range(6, 10)]
                        for enum, x in enumerate(zone_label):
                            df_excel.loc[df_lin_index, x] = df_list[0].iloc[enum + 6, 0]
                extended_report = [set_number]
                extended_report_win = pd.DataFrame(data=None, columns=['№'] + zone_label)
                extended_report_lin = pd.DataFrame(data=None, columns=['№'] + zone_label)
                if 'default' not in zone_check:
                    df_extended_report.loc[len(df_extended_report)] = extended_report +\
                                                                      [None]*(len(df_extended_report.columns) - 1)
                for element in zone_check:
                    if 'default' not in zone_check:
                        extended_report = [element]
                    replace_zone_check = []
                    for i in zone_check[element]:
                        if '<' in i:
                            replace_zone_check.append(0)
                        elif '-' in i:
                            replace_zone_check.append(-1)
                        else:
                            replace_zone_check.append(round(float(i), 1))
                    zone_check[element] = replace_zone_check
                    for enum, elem in enumerate(zone_check[element]):
                        extended_report.append(elem)
                        if elem > user_zone[enum]:
                            if element == 'Linux':
                                report[set_number]['error_lin'] = True
                            else:
                                report[set_number]['error_win'] = True
                    if 'default' in zone_check:
                        df_extended_report.loc[len(df_extended_report)] = extended_report
                        break
                    else:
                        if self.one_table:
                            df_extended_report.loc[len(df_extended_report)] = extended_report
                        elif element == 'Windows':
                            extended_report_win.loc[len(extended_report_win)] = extended_report
                        elif element == 'Linux':
                            extended_report_lin.loc[len(extended_report_lin)] = extended_report
                # Идём по первой таблице, сравниваем зоны (может в отдельную функцию)
                self.event.wait()
                if self.window_check.stop_threading:
                    raise CancelException()
                progress = progress + percent
                self.set_line_progress(f'Выполнено {int(progress)} %')
                self.progress_value.emit(int(progress))
                self.logging.info("Проверяем и красим ячейки в таблице")
                self.line_doing.emit(f'Проверяем зоны в {str(name_doc)} ({now_doc} из {all_doc})')
                # Столбцы для проверки (считать с 0 в Word)
                zone_col = [5, 7, 9, 11, 13] if self.department else [8, 9, 10, 11]

                def cur_mode_and_sys(find):
                    if '(' in find:
                        cm = re.findall(r"\(([^)]*)\)", find)[0]
                    else:
                        cm = find
                    if 'windows' not in find.lower() and 'linux' not in find.lower():
                        cns = ''
                    else:
                        cns = 'Windows' if 'windows' in find.lower() else 'Linux'
                        # cns = re.findall(r"\)([^()]*)\(", find)[0].strip()
                    if enum_df == 1:
                        cm = 'Power' if len(cm) > 10 else cm + '_pwr'
                    return cm, cns
                if self.one_table:
                    progress = progress + percent
                    self.set_line_progress(f'Выполнено {int(progress)} %')
                    self.progress_value.emit(int(progress))
                    continue
                for enum_df, df in enumerate([df_list[1], df_list[2]]):
                    if df.shape[0] < 2:
                        continue
                    # Проверка чтобы не бегать если третья таблица не мощность
                    if self.department is False and df.shape[1] > 12:
                        continue
                    if df.shape[1] < 10:
                        continue
                    find_st = df.loc[(df[0] == 'Максимальные значения') |
                                     (df.apply(lambda x: x[0] == x[zone_col[0]] == x[zone_col[-1]] ==
                                                                 x[zone_col[1]] == x[zone_col[-2]], axis=1))]
                    if self.department is False:
                        find_st = find_st.drop(0)
                    if self.department:
                        mag_line = [i for i in find_st.index.tolist() if 'магнитная составляющая' in find_st.loc[i, 0]]
                        find_st = find_st.drop(sorted(mag_line + [i + 1 for i in mag_line]))
                    # extend report
                    name_mode = df.drop([i for i in range(df.shape[1]) if i not in zone_col and i != 0], axis=1)
                    if name_mode.iloc[0, 0] == '2 категория':
                        name_mode = name_mode.drop(0)
                    name_mode = name_mode.loc[(df.apply(lambda x: x[0] == x[zone_col[0]] == x[zone_col[-1]] ==
                                                                  x[zone_col[1]] == x[zone_col[-2]]
                                                        and x[zone_col[0]] != 'Опасные сигналы не обнаружены', axis=1))]
                    # Тут продолжаем плохо дропнулось
                    if self.department:
                        mag_line = [i for i in name_mode.index.tolist() if 'магнитная составляющая'
                                    in name_mode.loc[i, 0]]
                        name_mode = name_mode.drop(mag_line)
                    name_mode[0] = list(map(cur_mode_and_sys, name_mode[0].to_numpy().tolist()))
                    extend_line = [e for e, x in enumerate(find_st.index.isin(name_mode.index.tolist())) if not x]
                    for col in zone_col:
                        list_val = find_st.iloc[extend_line, col].tolist()
                        name_mode[col] = list_val
                    for line in range(name_mode.shape[0]):
                        app_line = name_mode.iloc[line].tolist()
                        app_line[0] = app_line[0][0]
                        if self.win_lin:
                            if name_mode.iloc[line, 0][1] == 'Windows':
                                extended_report_win.loc[len(extended_report_win)] = app_line
                            else:
                                extended_report_lin.loc[len(extended_report_lin)] = app_line
                        else:
                            df_extended_report.loc[len(df_extended_report)] = app_line
                    find_zone = pd.DataFrame()
                    info_mode = 'ПЭМИ' if enum_df == 0 else 'ЦП'
                    self.line_doing.emit(f'Подсвечиваем ошибки в режиме {info_mode} {str(name_doc)}'
                                         f' ({now_doc} из {all_doc})')

                    def previous_val(lst, ind_):
                        if ind_ == 0:
                            return 0
                        else:
                            return ind_ - 1 if len(re.findall(r'[a-zA-Zа-яА-ЯёЁ-]', lst[ind_ - 1])) == 0 else\
                                previous_val(lst, ind_ - 1)

                    for index, zone in enumerate(zone_col):
                        start_val = df[zone].to_numpy().tolist()
                        all_zone = [start_val[previous_val(start_val, i)]
                                    if '-' in val and len(re.findall(r'[a-zA-Zа-яА-ЯёЁ]', val)) == 0 else val
                                    for i, val in enumerate(start_val)]
                        find_zone[zone] = list(map(lambda x: True if \
                            len(re.findall(r'[a-zA-Zа-яА-ЯёЁ]', x)) == 0 and float(user_zone[index]) < float(x) \
                            else False, all_zone))
                    find_zone = find_zone.drop(list(find_st.index.tolist()))
                    index_plus = 2 if self.department else 1
                    table_val = table_value if enum_df == 0 else table_pwr
                    first_line = []

                    def set_bg_color(tbl_cell):
                        tblCellProperties = tbl_cell.get_or_add_tcPr()
                        cl_shading = OxmlElement('w:shd')
                        cl_shading.set(qn('w:fill'), "FFFF00")
                        tblCellProperties.append(cl_shading)
                        return True

                    for zone in zone_col:
                        error_list = find_zone[zone].loc[find_zone[zone] == True].index.tolist()
                        first_line = first_line + error_list
                        [set_bg_color(table_val.rows[error + index_plus].cells[zone]._tc) for error in error_list]
                    if len(first_line) > 0:
                        first_line.sort()
                        first_line = [el for el, _ in groupby(first_line)]
                        ind_sys = find_st.index.tolist()
                        name_sys_ind = []
                        for fl in first_line:
                            for en, ind in enumerate(ind_sys):
                                if fl < ind:
                                    name_sys_ind.append(ind_sys[en - 1])
                                    break
                        if self.department:
                            frequency = (df.iloc[[int(x) for x in first_line], 0].tolist())
                        else:
                            frq_fstek = pd.Series((df[1].to_numpy() + '-' + df[2].to_numpy()).tolist())
                            frequency = frq_fstek.iloc[[int(x) for x in first_line]].tolist()
                        for enum_name, name_sys in enumerate(df.iloc[name_sys_ind, 0].tolist()):
                            cur_mode, cur_name_system = cur_mode_and_sys(name_sys)
                            sys_err = report[set_number]['lin'] if 'linux' in name_sys.lower()\
                                else report[set_number]['win']
                            table_val = table_value if enum_df == 0 else table_pwr
                            if cur_mode not in sys_err.keys():
                                sys_err[cur_mode] = []
                            sys_err[cur_mode].append(frequency[enum_name])
                        if self.department:
                            [set_bg_color(table_val.rows[error + index_plus].cells[0]._tc) for error in first_line]
                        else:
                            [set_bg_color(table_val.rows[error + index_plus].cells[1]._tc) for error in first_line]
                            [set_bg_color(table_val.rows[error + index_plus].cells[2]._tc) for error in first_line]
                df_extended_report = df_extended_report.append(extended_report_win)
                df_extended_report = df_extended_report.append(extended_report_lin)
                self.logging.info("Сохраняем документ")
                doc.save(pathlib.Path(self.path, name_doc))
                if self.window_check.stop_threading:
                    break
                progress = progress + percent
                self.set_line_progress(f'Выполнено {int(progress)} %')
                self.progress_value.emit(int(progress))
            if self.window_check.stop_threading:
                self.logging.warning('Прервано пользователем')
                os.chdir(self.default_path)
                self.status_finish.emit('zone_checked', str(self))
                time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
                self.window_check.close()
                return
            self.logging.info("Формируем excel")
            # Везде к строкам и колонкам прибавляем 2, потому что в excel начинается с 1 (+1)
            # а так же потому что в df заголовок идёт без индекса (+1)
            if pathlib.Path.exists(pathlib.Path(self.path, 'Зоны.xlsx')):
                while True:
                    try:
                        os.remove(pathlib.Path(self.path, 'Зоны.xlsx'))
                        break
                    except PermissionError:
                        self.info_value.emit('Вопрос?', 'Файл «Зоны.xlsx» в проверемой папке должен быть перезаписан. '
                                                        'При необходимости сохраните файл в другое место и закройте '
                                                        'его. После этого нажмите «Да» для продолжения или '
                                                        '«Нет» для прерывания')
                        self.event.clear()
                        self.event.wait()
                        if self.window_check.stop_threading:
                            self.logging.warning('Прервано пользователем')
                            os.chdir(self.default_path)
                            self.status_finish.emit('zone_checked', str(self))
                            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
                            self.window_check.close()
                            raise CancelException()
            if self.department:
                df_extended_report = df_extended_report[['№', 'Нос.', 'r1', 'r1`', 'Стац.', 'Воз.']
                                                        + [str(i) for i in range(df_extended_report.shape[1] - 6)]]
                df_extended_report.columns = ['№'] + zone_label + [str(i) for i in range(df_extended_report.shape[1]-6)]
            with pd.ExcelWriter(pathlib.Path(self.path, 'Зоны.xlsx'), engine='openpyxl', mode='w') as writer:
                df_excel.to_excel(writer, sheet_name='Зоны', index=False, header=False, startrow=1)
                if self.extend_report:
                    df_extended_report.to_excel(writer, sheet_name='Отчёт', index=False, header=False)
            thin = Side(border_style="thin", color="000000")
            wb = load_workbook(pathlib.Path(self.path, 'Зоны.xlsx'))
            ws = wb['Зоны']
            set_number = 0
            end_range = 7 if self.department else 6
            for enum, element in enumerate(zone_label):
                ws.cell(1, enum + 2).value = element
                ws.cell(1, enum + 2).font = Font(name="Times New Roman", size="11")
                ws.cell(1, enum + 2).alignment = Alignment(horizontal="center", vertical="center")
            append_err = False
            sys_err = True
            for row in range(2, df_excel.shape[0] + 2):
                for col in range(1, end_range):
                    if col > 1 and ws.cell(row, col).value is None:
                        append_err = True
                        break
                    if col == 1:
                        if ws.cell(row, col).value in report.keys():
                            set_number = ws.cell(row, col).value
                            round_set(ws.cell(row, col))
                        if (isinstance(ws.cell(row, col).value, str) and ('Windows' in ws.cell(row, col).value or
                                                                          'Linux' in ws.cell(row, col).value)):
                            sys_err = True if 'Windows' in ws.cell(row, col).value else False
                    if sys_err and report[set_number]['error_win']:
                        ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                             fill_type="solid")
                    elif sys_err is False and report[set_number]['error_lin']:
                        ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                             fill_type="solid")
                    else:
                        ws.cell(row, col).fill = PatternFill(start_color='92D050', end_color='92D050',
                                                             fill_type="solid")
                    ws.cell(row, col).font = Font(name="Times New Roman", size="11")
                    ws.cell(row, col).alignment = Alignment(horizontal="center", vertical="center")
                    if col > 1:
                        if ws.cell(row, col).value == 0 or ws.cell(row, col).value == '0':
                            ws.cell(row, col).value = '<0.1'
                        elif ws.cell(row, col).value == '-':
                            ws.cell(row, col).value = '-'
                        elif '<' not in ws.cell(row, col).value:
                            if float(ws.cell(row, col).value) > float(self.zone[col - 2]):
                                round_number(ws.cell(row, col), bold=True, align=True)
                            else:
                                round_number(ws.cell(row, col), bold=False, align=True)
                        ws.cell(row, col).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                if report[set_number]['error_win'] or report[set_number]['error_lin']:
                    if append_err:
                        append_err = False
                        continue
                    mode_frq = report[set_number]['win'] if sys_err else report[set_number]['lin']
                    errors_col = end_range
                    for key in mode_frq:
                        ws.cell(row, errors_col).value = key
                        ws.cell(row, errors_col).font = Font(bold=True, name="Times New Roman", size="11")
                        errors_col += 1
                        for value in mode_frq[key]:
                            round_number(ws.cell(row, errors_col), val=value, string=True if '-' in value else False)
                            errors_col += 1
            if self.extend_report:
                ws = wb['Отчёт']
                set_row = 0
                green = True
                for row in range(1, df_extended_report.shape[0] + 1):
                    for col in range(1, df_extended_report.shape[1] + 1):
                        if col > 1 and ws.cell(row, col).value is None:
                            break
                        ws.cell(row, col).font = Font(name="Times New Roman", size="11")
                        if col == 1:
                            if ws.cell(row, col).value in report.keys():
                                set_number = ws.cell(row, col).value
                                round_set(ws.cell(row, col))
                                set_row = row
                                ws.cell(row, col).fill = PatternFill(start_color='92D050', end_color='92D050',
                                                                     fill_type="solid")
                            # Проверка для покраски - если 1 система, то индекс строки номера комплекта совпадает
                            # и второе значение в строке не пустое.
                            # Если 2 системы, смотрим, чтобы тип значения был строкой (отфоматировали в предыдущем шаге)
                            # и содержалось название системы.
                            if row == set_row and ws.cell(row, 2).value is not None:
                                if report[set_number]['error_win']:
                                    ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                         fill_type="solid")
                                    sys_err = True
                                    green = False
                                else:
                                    ws.cell(row, col).fill = PatternFill(start_color='92D050', end_color='92D050',
                                                                         fill_type="solid")
                                    green = True

                            elif isinstance(ws.cell(row, col).value, str) and ('Windows' in ws.cell(row, col).value
                                                                               or 'Linux' in ws.cell(row, col).value):
                                if report[set_number]['error_win'] and 'Windows' in ws.cell(row, col).value:
                                    ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                         fill_type="solid")
                                    ws.cell(set_row, 1).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                           fill_type="solid")
                                    sys_err = True
                                    green = False
                                elif report[set_number]['error_lin'] and 'Linux' in ws.cell(row, col).value:
                                    ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                         fill_type="solid")
                                    ws.cell(set_row, 1).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                           fill_type="solid")
                                    sys_err = False
                                    green = False
                                else:
                                    ws.cell(row, col).fill = PatternFill(start_color='92D050', end_color='92D050',
                                                                         fill_type="solid")
                                    green = True
                                set_row = row
                        if col == 2 and isinstance(ws.cell(row, col).value, str) \
                                and 'не обнаружен' in ws.cell(row, col).value:
                            ws.cell(row, col).alignment = Alignment(horizontal="center", vertical="center")
                            for j in range(2, end_range):
                                ws.cell(row, j).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=end_range - 1)
                            break
                        if 1 < col < end_range:
                            ws.cell(row, col).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                            ws.cell(row, col).alignment = Alignment(horizontal="center", vertical="center")
                            if ws.cell(row, col).value == 0 or ws.cell(row, col).value == '0':
                                ws.cell(row, col).value = '<0.1'
                            elif ws.cell(row, col).value == '-':
                                ws.cell(row, col).value = '-'
                            if row == set_row:
                                if green:
                                    ws.cell(row, col).fill = PatternFill(start_color='92D050', end_color='92D050',
                                                                         fill_type="solid")
                                else:
                                    ws.cell(row, col).fill = PatternFill(start_color='FFC000', end_color='FFC000',
                                                                         fill_type="solid")
                            if (isinstance(ws.cell(row, col).value, float) or\
                                isinstance(ws.cell(row, col).value, int) or\
                                    all([True for i in ws.cell(row, col).value
                                         if i in ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9', '.']])) and\
                                    ws.cell(row, col).value != '<0.1' and ws.cell(row, col).value != '-':
                                round_number(ws.cell(row, col), bold=False, align=True, string=False)
                                if float(ws.cell(row, col).value) > float(self.zone[col - 2]):
                                    round_number(ws.cell(row, col), bold=True, align=True, string=False)
                    mode_frq = report[set_number]['win'] if sys_err else report[set_number]['lin']
                    if ws.cell(row, 1).value in mode_frq.keys():
                        errors_col = end_range
                        for value in mode_frq[ws.cell(row, 1).value]:
                            round_number(ws.cell(row, errors_col), val=value, string=True if '-' in value else False)
                            errors_col += 1
            wb.save(pathlib.Path(self.path, 'Зоны.xlsx'))
            os.startfile(pathlib.Path(self.path, 'Зоны.xlsx'))
            self.logging.info(f"Проверка зон в папке «{self.name_dir}» успешно завершена")
            self.progress_value.emit(int(100))
            os.chdir(self.default_path)
            self.status.emit(f"Проверка зон в папке «{self.name_dir}» успешно завершена")
            self.status_finish.emit('zone_checked', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            # print(datetime.datetime.now() - start_time)
            return
        except CancelException:
            self.logging.warning(f"Проверка зон в папке «{self.name_dir}» отменена пользователем")
            self.status.emit(f"Проверка зон в папке «{self.name_dir}» отменена пользователем")
            os.chdir(self.default_path)
            self.status_finish.emit('zone_checked', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
        except BaseException as es:
            self.logging.error(es)
            self.logging.error(traceback.format_exc())
            self.logging.warning(f"Проверка зон в папке «{self.name_dir}» не заврешена из-за ошибки")
            self.info_value.emit('УПС!', 'Работа программы завершена из-за непредвиденной ошибки')
            self.event.clear()
            self.event.wait()
            self.status.emit(f"Ошибка при проверке зон в папке «{self.name_dir}»")
            os.chdir(self.default_path)
            self.status_finish.emit('zone_checked', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
