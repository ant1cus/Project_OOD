# import datetime
import copy
import shutil
import zipfile

from docx.shared import Pt, Cm
from lxml import etree
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pathlib
import threading
import time

import docx
import re
import traceback
from PyQt5.QtCore import QThread, pyqtSignal
from DoingWindow import CheckWindow


class CancelException(Exception):
    pass


class DeleteHeaderFooter(QThread):
    status_finish = pyqtSignal(str, str)
    progress_value = pyqtSignal(int)
    info_value = pyqtSignal(str, str)
    status = pyqtSignal(str)
    line_progress = pyqtSignal(str)
    line_doing = pyqtSignal(str)

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path = incoming_data['path']
        self.conclusion = incoming_data['conclusion']
        self.protocol = incoming_data['protocol']
        self.prescription = incoming_data['prescription']
        self.project_prescription = incoming_data['project_prescription']
        self.old_director = incoming_data['old_director']
        self.new_director = incoming_data['new_director']
        self.margin = incoming_data['margin']
        self.margin_left = incoming_data['margin_left']
        self.margin_top = incoming_data['margin_top']
        self.margin_right = incoming_data['margin_right']
        self.margin_bottom = incoming_data['margin_bottom']
        self.logging = incoming_data['logging']
        self.queue = incoming_data['queue']
        self.default_path = incoming_data['default_path']
        self.all_doc = 0
        self.now_doc = 0
        self.event = threading.Event()
        self.event.set()
        self.move = incoming_data['move']
        self.name_dir = pathlib.Path(self.path).name
        title = f'Удаление колонтитулов в папке «{self.name_dir}»'
        self.window_check = CheckWindow(self.default_path, self.event, self.move, title)
        self.progress_value.connect(self.window_check.progressBar.setValue)
        self.line_progress.connect(self.window_check.lineEdit_progress.setText)
        self.line_doing.connect(self.window_check.lineEdit_doing.setText)
        self.info_value.connect(self.window_check.info_message)
        self.window_check.show()

    def run(self):
        try:
            current_progress = 0
            self.logging.info('Начинаем обезличивать документы')
            self.progress_value.emit(int(current_progress))
            self.line_progress.emit(f'Выполнено {int(current_progress)} %')
            files = [file for file in os.listdir(self.path) if '~' not in file and file.endswith('.docx')]
            percent = 100 / len(files)
            self.all_doc = len(files)
            no_header = []
            for element in files:
                self.event.wait()
                if self.window_check.stop_threading:
                    raise CancelException()
                self.now_doc += 1
                self.logging.info(f'Обезличиваем документ {element}')
                self.line_doing.emit(f'Обезличиваем документ {element} ({self.now_doc} из {self.all_doc})')
                if 'приложение' in element.lower():
                    self.logging.info('Работаем с ворд документом')
                    doc = docx.Document(pathlib.Path(self.path, element))  # Открываем
                    sub = False
                    for paragraph in doc.paragraphs:
                        if 'к протоколу' in paragraph.text:
                            paragraph.text = re.sub('к протоколу', 'к выписке из протокола', paragraph.text)
                            for run in paragraph.runs:
                                run.font.bold = False
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                            sub = True
                            break
                        if 'к выписке из протокола' in paragraph.text:
                            break
                    if sub:
                        doc.save(pathlib.Path(self.path, element))
                    else:
                        no_header.append(element)
                        self.logging.info(f'Документ {element} скорее всего документ уже обезличен.'
                                          f' Переходим к следующему')
                    current_progress += percent
                    self.progress_value.emit(int(current_progress))
                    self.line_progress.emit(f'Выполнено {int(current_progress)} %')
                    continue
                temp_docx = os.path.join(self.path + '\\', element)
                temp_zip = os.path.join(self.path + '\\', element + ".zip")
                temp_folder = os.path.join(self.path + '\\', "template")
                if os.path.exists(temp_zip):
                    shutil.rmtree(temp_zip)
                if os.path.exists(temp_folder):
                    shutil.rmtree(temp_folder)
                if os.path.exists(self.path + '\\zip'):
                    shutil.rmtree(self.path + '\\zip')
                os.rename(temp_docx, temp_zip)
                os.mkdir(self.path + '\\zip')
                self.logging.info('Извлекаем архив')
                with zipfile.ZipFile(temp_zip) as my_document:
                    my_document.extractall(temp_folder)
                pages_xml = os.path.join(temp_folder, "word", "document.xml")
                tree = etree.parse(pages_xml)
                root = tree.getroot()
                self.logging.info('Получаем тэг для копирования')
                header = ''

                ############
                def req_for_tag(tag_, header_):
                    for enum_, i_ in enumerate(tag_):
                        keys = [key for key in i_.attrib.keys() if '}id' in key]
                        for key in keys:
                            if 'rId8' in i_.attrib[key] or 'rId9' in i_.attrib[key] or 'rId10' in i_.attrib[key]:
                                header_ = copy.deepcopy(tag_)
                                return header_
                        if len(i_):
                            header_ = req_for_tag(i_, header_)
                            if len(header_):
                                return header_
                    return header_

                for i in range(len(root[0]) - 1, 0, -1):
                    header = req_for_tag(root[0][i], header)
                    if len(header):
                        break
                ############
                self.logging.info('Собираем и удаляем архив')
                os.remove(temp_zip)
                shutil.make_archive(temp_zip.replace(".zip", ""), 'zip', temp_folder)
                os.rename(temp_zip, temp_docx)  # rename zip file to docx
                while True:
                    try:
                        shutil.rmtree(temp_folder)
                        shutil.rmtree(self.path + '\\zip')
                        break
                    except OSError as es:
                        self.logging.error(es)
                        self.logging.error(traceback.format_exc())
                        self.logging.info('Ошибка с удалением, пробуем ещё раз')
                self.logging.info('Работаем с ворд документом')
                doc = docx.Document(pathlib.Path(self.path, element))  # Открываем

                def paragraph_del(par):
                    par.text = None
                    paragraph_ = par._element
                    paragraph_.getparent().remove(paragraph_)
                    paragraph_._p = paragraph_._element = None

                list_paragraph = []
                for enum, paragraph in enumerate(doc.sections[0].first_page_header.paragraphs):
                    if 'экз.' in paragraph.text.lower():
                        list_paragraph = [i for i in range(enum + 1)]
                        break
                for paragraph in list_paragraph:
                    doc.sections[0].first_page_header.paragraphs[paragraph].text = None
                for paragraph in range(len(list_paragraph) - 1):
                    p = doc.sections[0].first_page_header.paragraphs[paragraph]._element
                    p.getparent().remove(p)
                    p._p = p._element = None
                p = doc.sections[0].first_page_header.paragraphs[0]._element
                p.getparent().remove(p)
                p._p = p._element = None
                number = doc.sections[0].first_page_footer.paragraphs[0].text
                doc.sections[0].first_page_footer.paragraphs[0].text = None
                date = []
                if len(doc.sections) == 1:
                    for paragraph in doc.sections[len(doc.sections) - 1].footer.paragraphs:
                        if re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text):
                            date = re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text)
                        if 'Б/ч' in paragraph.text:
                            paragraph_del(paragraph)
                            break
                        paragraph_del(paragraph)
                else:
                    doc.sections[0].footer.paragraphs[0].text = None
                    if doc.sections[len(doc.sections) - 1].different_first_page_header_footer:
                        for paragraph in doc.sections[len(doc.sections) - 1].first_page_footer.paragraphs:
                            if re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text):
                                date = re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text)
                            if 'Б/ч' in paragraph.text:
                                paragraph_del(paragraph)
                                break
                            paragraph_del(paragraph)
                        doc.sections[len(doc.sections) - 1].first_page_header.is_linked_to_previous = False  # header
                        doc.sections[len(doc.sections) - 1].first_page_footer.is_linked_to_previous = False  # Футер
                    else:
                        for paragraph in doc.sections[len(doc.sections) - 1].footer.paragraphs:
                            if re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text):
                                date = re.findall(r'\d{2}\.\d{2}\.\d{4}', paragraph.text)
                            if 'Б/ч' in paragraph.text:
                                paragraph_del(paragraph)
                                break
                            paragraph_del(paragraph)
                if len(date) == 0:
                    no_header.append(element)
                    self.logging.info(f'В документе {element} не нашлась дата в колонтитуле, скорее всего документ уже '
                                      f'обезличен. Переходим к следующему')
                    current_progress += percent
                    self.progress_value.emit(int(current_progress))
                    self.line_progress.emit(f'Выполнено {int(current_progress)} %')
                    continue
                while True:
                    flag_for_exit = 0
                    if flag_for_exit == 3:
                        break
                    try:
                        doc.save(pathlib.Path(self.path, element))
                        break
                    except PermissionError:
                        self.logging.warning(f'Документ {element} не удалось сохранить с первого раза, ждём...')
                        flag_for_exit += 1
                        time.sleep(3)
                doc = docx.Document(pathlib.Path(self.path, element))  # Открываем
                self.logging.info('Удаляем лишние параграфы в конце документа')
                sectPrs = doc._element.xpath(".//w:pPr/w:sectPr")
                for sectPr in sectPrs:
                    sectPr.getparent().remove(sectPr)
                doc.save(pathlib.Path(self.path, element))
                doc = docx.Document(pathlib.Path(self.path, element))  # Открываем
                self.logging.info('Переименовываем заголовок')
                name = element.rpartition('.')[0]
                size = 12

                def size_pt(search_paragraph):
                    if search_paragraph.runs[0].font.size is None:
                        find_size = search_paragraph.style.font.size.pt
                    else:
                        find_size = search_paragraph.runs[0].font.size.pt
                    return find_size

                if 'Заключение СП' in name:
                    name = re.sub('Заключение СП', 'Заключение', name)
                for i, paragraph in enumerate(doc.paragraphs):
                    if self.old_director:
                        if re.findall(self.old_director, paragraph.text):
                            paragraph.text = re.sub(self.old_director, self.new_director, paragraph.text)
                            size = size_pt(paragraph)
                            for run in paragraph.runs:
                                run.font.bold = False
                                run.font.size = Pt(size)
                    if re.findall(r'«\d{2}»\s\w+\s\d{4}\s[г]\.', paragraph.text):
                        size = size_pt(paragraph)
                        paragraph.text = re.sub(r'«\d{2}»\s\w+\s\d{4}\s[г]\.', 'date', paragraph.text)
                        for run in paragraph.runs:
                            run.font.size = Pt(size)
                    if re.findall(name.partition(' ')[0].upper(), paragraph.text) or\
                            re.findall(name.partition(' ')[0][:-1].upper(), paragraph.text):
                        size = size_pt(paragraph)
                        name_doc = name.partition(' ')[0] if name.partition(' ')[0] else name.partition('.')[0]
                        if name_doc.lower() == 'протокол' or name_doc.lower() == 'акт':
                            name_doc += 'а'
                        else:
                            name_doc = list(name_doc)
                            name_doc[-1] = 'я'
                            name_doc = ''.join(name_doc)
                        if re.findall(r'ПРОЕКТ ПРЕДПИСАНИЯ', paragraph.text):
                            if self.project_prescription:
                                paragraph.text = re.sub('ПРОЕКТ ПРЕДПИСАНИЯ',
                                                        'ВЫПИСКА ИЗ ПРОЕКТА ПРЕДПИСАНИЯ',
                                                        paragraph.text)
                            else:
                                paragraph.text = re.sub('ПРОЕКТ ПРЕДПИСАНИЯ',
                                                        'ВЫПИСКА ИЗ ПРЕДПИСАНИЯ',
                                                        paragraph.text)
                        else:
                            paragraph.text = re.sub(name.partition(' ')[0].upper(),
                                                    'ВЫПИСКА ИЗ ' + name_doc.upper(),
                                                    paragraph.text)
                        # paragraph.text = re.sub(name.partition(' ')[0].upper(),
                        #                         'ВЫПИСКА ИЗ ' + name_doc.upper(),
                        #                         paragraph.text)
                        for run in paragraph.runs:
                            run.font.bold = False
                            run.font.name = 'Times New Roman'
                        doc.paragraphs[i + 1].insert_paragraph_before('Уч. № ' + number + ' от ' + date[0])
                        for par_format in range(4):
                            doc.paragraphs[i + par_format].paragraph_format.first_line_indent = None
                            doc.paragraphs[i + par_format].paragraph_format.left_indent = Cm(0)
                        for j in [i, i + 1]:
                            doc.paragraphs[j].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in doc.paragraphs[j].runs:
                                run.font.bold = True
                                run.font.size = Pt(size)
                        break
                doc.save(pathlib.Path(self.path, element))
                doc = docx.Document(pathlib.Path(self.path, element))  # Открываем
                if 'Заключение' in name or 'Акт' in name:
                    executor = self.conclusion
                elif 'Протокол' in name:
                    executor = self.protocol
                else:
                    executor = self.prescription
                doc.add_paragraph('Выписка верна\n\n' + executor)
                doc.paragraphs[len(doc.paragraphs) - 1].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравниваем
                for run in doc.paragraphs[len(doc.paragraphs) - 1].runs:
                    run.font.size = Pt(size)
                    run.font.name = 'Times New Roman'
                doc.save(pathlib.Path(self.path, element))
                self.logging.info('Извлекаем архив и добавляем шапку')
                while True:
                    rename_chance = 1
                    if rename_chance == 4:
                        break
                        # os.rename(temp_docx, temp_zip)
                        # break
                    try:
                        os.rename(temp_docx, temp_zip)
                        break
                    except BaseException:
                        self.logging.warning(f'Не смог переименовать файл {rename_chance} раз, ждём...')
                        rename_chance += 1
                        time.sleep(3)
                os.mkdir(self.path + '\\zip')
                with zipfile.ZipFile(temp_zip) as my_document:
                    my_document.extractall(temp_folder)
                if len(header):
                    pages_xml = os.path.join(temp_folder, "word", "document.xml")
                    tree = etree.parse(pages_xml)
                    root = tree.getroot()
                    for i, e in reversed(list(enumerate(root[0]))):
                        if e.tag.rpartition('}')[2] == 'sectPr':
                            e.getparent().replace(e, header)
                            break
                    tree.write(os.path.join(temp_folder, "word", "document.xml"), encoding="UTF-8",
                               xml_declaration=True)
                self.logging.info('Удаляем архив')
                os.remove(temp_zip)
                while True:
                    rename_chance = 1
                    if rename_chance == 4:
                        break
                    try:
                        shutil.make_archive(temp_zip.replace(".zip", ""), 'zip', temp_folder)
                        break
                    except BaseException:
                        self.logging.warning(f'Не смог сделать архив {rename_chance} раз, ждём...')
                        rename_chance += 1
                        time.sleep(3)
                # shutil.make_archive(temp_zip.replace(".zip", ""), 'zip', temp_folder)
                os.rename(temp_zip, temp_docx)  # rename zip file to docx
                while True:
                    try:
                        shutil.rmtree(temp_folder)
                        shutil.rmtree(self.path + '\\zip')
                        break
                    except BaseException as es:
                        self.logging.error(es)
                        self.logging.error(traceback.format_exc())
                        self.logging.info('Ошибка с удалением, пробуем ещё раз')
                if self.margin:
                    self.logging.info('Меняем поля в документе')
                    doc = docx.Document(pathlib.Path(self.path, element))  # Открываем
                    for section in doc.sections:
                        section.left_margin = Cm(self.margin_left)
                        section.top_margin = Cm(self.margin_top)
                        section.right_margin = Cm(self.margin_right)
                        section.bottom_margin = Cm(self.margin_bottom)
                    doc.save(pathlib.Path(self.path, element))
                current_progress += percent
                self.progress_value.emit(int(current_progress))
                self.line_progress.emit(f'Выполнено {int(current_progress)} %')
            if len(no_header):
                self.logging.info("Выводим файлы без колонтитулов")
                self.info_value.emit('Внимание!', 'Следующие документы уже обезличены:\n' + ', '.join(no_header))
                self.event.clear()
                self.event.wait()
            self.logging.info(f"Обезличивание документов в папке «{self.name_dir}» успешно завершено")
            self.progress_value.emit(100)
            os.chdir(self.default_path)
            self.status.emit(f"Обезличивание документов в папке «{self.name_dir}» успешно завершено")
            self.status_finish.emit('delete_header_footer', str(self))
            time.sleep(1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            # print(datetime.datetime.now() - start_time)
            return
        except CancelException:
            self.logging.warning(f"Обезличивание документов в папке «{self.name_dir}» отменено пользователем")
            self.status.emit(f"Обезличивание документов в папке «{self.name_dir}» отменено пользователем")
            os.chdir(self.default_path)
            self.status_finish.emit('delete_header_footer', str(self))
            time.sleep(1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
        except BaseException as es:
            self.logging.error(es)
            self.logging.error(traceback.format_exc())
            self.logging.warning(f"Обезличивание документов в папке «{self.name_dir}» не завершено из-за ошибки")
            self.info_value.emit('УПС!', 'Работа программы завершена из-за непредвиденной ошибки')
            self.event.clear()
            self.event.wait()
            self.status.emit(f"Ошибка при обезличивании документов в папке «{self.name_dir}»")
            os.chdir(self.default_path)
            self.status_finish.emit('delete_header_footer', str(self))
            time.sleep(1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
