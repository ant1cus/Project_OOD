import os
import pathlib
import time

import docx
import pandas as pd
import xlwings as xw
from decimal import Decimal
import threading
import traceback
import pythoncom
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from DoingWindow import CheckWindow
from PyQt5.QtCore import QThread, pyqtSignal


class CancelException(Exception):
    pass


class LFGeneration(QThread):
    status_finish = pyqtSignal(str, str)
    progress_value = pyqtSignal(int)
    info_value = pyqtSignal(str, str)
    status = pyqtSignal(str)
    line_progress = pyqtSignal(str)
    line_doing = pyqtSignal(str)

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path_old = incoming_data['source']
        self.path_new = incoming_data['output']
        self.excel = incoming_data['excel']
        self.logging = incoming_data['logging']
        self.queue = incoming_data['queue']
        self.default_path = incoming_data['default_path']
        self.event = threading.Event()
        self.event.set()
        self.all_doc = 0
        self.now_doc = 0
        self.percent_progress = 0
        self.move = incoming_data['move']
        self.name_dir = pathlib.Path(self.path_new).name
        title = f'Генерация НЧ в папке «{self.name_dir}»'
        self.window_check = CheckWindow(self.default_path, self.event, self.move, title)
        self.progress_value.connect(self.window_check.progressBar.setValue)
        self.line_progress.connect(self.window_check.lineEdit_progress.setText)
        self.line_doing.connect(self.window_check.lineEdit_doing.setText)
        self.info_value.connect(self.window_check.info_message)
        self.window_check.show()

    def run(self):
        pythoncom.CoInitialize()
        current_progress = 0
        self.percent_progress = 100 / len(os.listdir(self.path_old))
        self.all_doc = len(os.listdir(self.path_old))
        self.logging.info('Начинаем генерацию НЧ')
        self.line_progress.emit(f'Выполнено {int(current_progress)} %')
        self.progress_value.emit(0)
        try:
            for element in os.listdir(self.path_old):
                self.now_doc += 1
                self.line_doing.emit(f'Вставляем данные в файл {element} ({self.now_doc} из {self.all_doc})')
                self.event.wait()
                if self.window_check.stop_threading:
                    raise CancelException()
                self.logging.info(f'Вставляем данные в файл {element}')
                self.logging.info('Открываем и закрываем excel для смены чисел')
                excel_app = xw.App(visible=False)
                excel_book = excel_app.books.open(self.excel)
                excel_book.save()
                excel_book.close()
                excel_app.quit()
                self.logging.info('Читаем таблицы')
                table_1 = pd.read_excel(self.excel, sheet_name='Таблица для вставки')
                table_2 = pd.read_excel(self.excel, sheet_name='Таблица для вставки с R2_r1_r1')
                self.logging.info('Открываем docx')
                doc = docx.Document(self.path_old + '\\' + element)
                table = doc.tables
                self.logging.info('Заполняем таблицу ' + str(len(table) - 3))
                for i, row in enumerate(table[len(table) - 3].rows):
                    if i != 0 and i != 1:
                        len_row = len(row.cells)
                        for j, cell in enumerate(row.cells):
                            self.logging.info(cell.text)
                            if j == 0:
                                name_device = cell.text.strip()
                                name_col = table_1.columns
                                data_row = table_1[table_1[name_col[1]] == name_device].index.to_list()[0]
                            elif 1 <= j < 4:
                                cell.text = str(Decimal(table_1.iloc[data_row, j + 1]).quantize(Decimal('1.0')))
                            elif j == 4:
                                cell.text = str(int(table_1.iloc[data_row, j + 1]))
                            elif j >= 5:
                                if len_row == 6:
                                    cell.text = str(Decimal(table_1.iloc[data_row, j + 2]).quantize(Decimal('1.0')))
                                    cell.paragraphs[
                                        0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
                                    break
                                else:
                                    cell.text = str(Decimal(table_1.iloc[data_row, j + 1]).quantize(Decimal('1.0')))
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
                self.logging.info('Заполняем таблицу ' + str(len(table) - 2))
                for i, row in enumerate(table[len(table) - 2].rows):
                    data_row += 1
                    if i == 0:
                        flag = 1 if len(row.cells) == 2 else 0
                    else:
                        for j, cell in enumerate(row.cells):
                            if 'R2' not in cell.text and 'r1' not in cell.text and len(cell.text) != 0:
                                name_device = cell.text.strip()
                                name_col = table_2.columns
                                data_row = table_2[table_2[name_col[0]] == name_device].index.to_list()[0]
                            else:
                                if j > 0:
                                    if type(table_2.iloc[data_row, j + flag]) == str:
                                        cell.text = table_2.iloc[data_row, j + flag]
                                    else:
                                        cell.text = str(Decimal(table_2.iloc[data_row,
                                                                             j + flag]).quantize(Decimal('1.0')))
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
                self.logging.info('Сохраняем документ')
                doc.save(pathlib.Path(self.path_new, element))
                current_progress += self.percent_progress
                self.line_progress.emit(f'Выполнено {int(current_progress)} %')
                self.progress_value.emit(int(current_progress))
            self.line_progress.emit(f'Выполнено 100 %')
            self.logging.info(f"Генерация НЧ в папке «{self.name_dir}» успешно завершена")
            self.progress_value.emit(int(100))
            os.chdir(self.default_path)
            self.status.emit(f"Генерация НЧ в папке «{self.name_dir}» успешно завершена")
            self.status_finish.emit('generate_lf', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            # print(datetime.datetime.now() - start_time)
            return
        except CancelException:
            os.chdir(self.default_path)
            self.status_finish.emit('generate_lf', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            self.logging.warning(f"Генерация НЧ в папке «{self.name_dir}» отменена пользователем")
            self.status.emit(f"Генерация НЧ в папке «{self.name_dir}» отменена пользователем")
            return
        except BaseException as es:
            self.logging.error(es)
            self.logging.error(traceback.format_exc())
            self.logging.warning(f"Генерация НЧ в папке «{self.name_dir}» не заврешена из-за ошибки")
            self.info_value.emit('УПС!', 'Работа программы завершена из-за непредвиденной ошибки')
            self.event.clear()
            self.event.wait()
            self.status.emit(f"Ошибка при генерации НЧ в папке «{self.name_dir}»")
            os.chdir(self.default_path)
            self.status_finish.emit('generate_lf', str(self))
            time.sleep(0.1)  # Не удалять, не успевает отработать emit status_finish. Может потом
            self.window_check.close()
            return
