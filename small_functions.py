import re
from pathlib import Path

import pandas as pd
from docx import Document
import json


def rewrite(path, data, widget=False, visible=False, order=False):
    if widget:
        dict_load = data
    else:
        name = visible if visible else order
        with open(Path(path, 'Настройки.txt'), "r", encoding='utf-8-sig') as f:  # Открываем
            dict_load = json.load(f)  # Загружаем данные
            dict_load['gui_settings'][name] = data
    with open(Path(path, 'Настройки.txt'), 'w', encoding='utf-8-sig') as f:  # Пишем в файл
        json.dump(dict_load, f, ensure_ascii=False, sort_keys=True, indent=4)


def read_description_file(path: Path) -> list:
    mode = []
    try:
        for file in Path(path).glob('*.txt'):
            try:
                with open(file, mode='r', encoding="utf-8-sig") as f:
                    mode_1 = f.readlines()
                    mode_1 = [line.rstrip() for line in mode_1]
            except UnicodeDecodeError:
                with open(file, mode='r') as f:
                    mode_1 = f.readlines()
                    mode_1 = [line.rstrip() for line in mode_1]
            mode = [x for x in mode_1 if x]
    except BaseException:
        pass
    return mode


def read_mode_docx(path: Path, win_lin: bool, department: bool) -> pd.DataFrame:
    mode = pd.DataFrame()
    try:
        for file in Path(path).glob('*.docx'):
            doc = Document(str(Path(path, file)))
            table_value = doc.tables[1] if win_lin and department else doc.tables[2]
            table_pwr = doc.tables[3] if win_lin and department else doc.tables[4]
            df_list = []
            for table_val in [table_value, table_pwr]:
                df = pd.DataFrame()
                for index in range(len(table_val.columns)):
                    df[index] = list(map(lambda val: '0' if '<' in val.text else val.text.replace(',', '.'),
                                         table_val.column_cells(index)))
                df_list.append(df)
            for i in [0, 1]:
                cat_3 = df_list[i].loc[df_list[i][0] == '3 категория'].index.to_list()
                if cat_3 and cat_3[0]:
                    df_list[i] = df_list[i].drop(index=[i for i in range(cat_3[0], df_list[i].shape[0])])

            def cur_mode_and_sys(find):
                if '(' in find:
                    cm = re.findall(r"\(([^)]*)\)", find)[0]
                else:
                    cm = find
                if 'windows' not in find.lower() and 'linux' not in find.lower():
                    cns = ''
                else:
                    cns = 'Windows' if 'windows' in find.lower() else 'Linux'
                if i == 1:
                    cm = 'Power' if len(cm) > 10 else cm + '_pwr'
                return cm, cns
            for i in range(2):
                df_list[i] = df_list[i].drop(index=[0, 1])
                df_list[i].reset_index(drop=True, inplace=True)
                df_list[i] = df_list[i][df_list[i].eq(df_list[i].iloc[:, 0], axis=0).all(axis=1)]
                df_list[i] = df_list[i][~(df_list[i] == 'Опасные сигналы не обнаружены').all(axis=1)]
                if department:
                    mag_line = [i for i in df.index.tolist() if 'магнитная составляющая' in df.loc[i, 0]]
                    df_list[i] = df_list[i].drop(mag_line)
                df_list[i][0] = list(map(cur_mode_and_sys, df_list[i][0].to_numpy().tolist()))
            mode = pd.concat([df_list[0][[0, 1]], df_list[1][[0, 1]]], ignore_index=True)
            mode.columns = ['mode_and_sys', 'name']
            mode['in_report'] = True
            break
    except BaseException:
        mode = pd.DataFrame()
    return mode
